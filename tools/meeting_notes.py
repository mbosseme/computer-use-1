"""
Daily Meeting Notes Automation (Outlook -> OneDrive Word doc + Otter ZIP append)
"""
import argparse
import datetime
import hashlib
import json
import logging
import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
try:
    from zoneinfo import ZoneInfo
except ImportError:
    from backports.zoneinfo import ZoneInfo

import docx
from dateutil import parser as date_parser

from agent_tools.graph.auth import GraphAuthenticator
from agent_tools.graph.client import GraphAPIClient, GraphClientConfig
from agent_tools.graph.env import load_graph_env
from agent_tools.llm.azure_openai_responses import AzureOpenAIResponsesClient
from agent_tools.llm.smoketest import _resolve_azure_config

# Default Config
ONEDRIVE_ROOT = Path("/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc")
NOTES_ROOT = ONEDRIVE_ROOT / "Notetaking" / "Meeting Notes"
OTTER_INBOX_DIR = ONEDRIVE_ROOT / "Otter Exports"
INDEX_DIR = NOTES_ROOT / "_index"
EVENTS_INDEX_FILE = INDEX_DIR / "events-index.json"
TRANSCRIPTS_INDEX_FILE = INDEX_DIR / "transcripts-index.json"

TZ_NAME = "America/New_York"
WIN_TZ_NAME = "Eastern Standard Time"   # for Graph Prefer header if needed

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# --- State Management ---
def load_index(path: Path) -> Dict[str, Any]:
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_index(path: Path, data: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)

# --- Filename + Helpers ---
def sanitize_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name[:50] # keep it short

def get_event_hash(event_id: str) -> str:
    return hashlib.sha1(event_id.encode('utf-8')).hexdigest()[:8]

def extract_attendees(attendees_list: List[Dict[str, Any]]) -> List[str]:
    names = []
    for a in attendees_list:
        email_addr = a.get("emailAddress", {})
        name = email_addr.get("name") or email_addr.get("address", "")
        # Get just last name if possible, or first
        parts = name.split()
        if len(parts) > 1:
            names.append(parts[-1])
        elif name:
            names.append(name)
    return names[:3] # keep it short

def build_filename(dt_start: datetime.datetime, subject: str, attendees: List[Dict[str, Any]], event_id: str) -> str:
    date_str = dt_start.strftime("%Y-%m-%d")
    time_str = dt_start.strftime("%H%M")
    safe_subj = sanitize_filename(subject) or "Untitled"
    attn_names = extract_attendees(attendees)
    attn_str = "-".join(attn_names) if attn_names else "NoAttendees"
    eid_hash = get_event_hash(event_id)
    
    fname = f"{date_str} {time_str} - {safe_subj} - {attn_str} - eid{eid_hash}.docx"
    # Keep within limits (180 chars)
    if len(fname) > 180:
        fname = f"{date_str} {time_str} - {safe_subj[:30]} - {eid_hash}.docx"
    return fname

# --- Graph Call ---
def get_graph_client(repo_root: Path) -> GraphAPIClient:
    env = load_graph_env(repo_root)
    authenticator = GraphAuthenticator(
        repo_root=repo_root,
        env=env
    )
    # Patch planner timezone if we want
    config = GraphClientConfig(
        base_url=env.base_url, 
        scopes=env.scopes, 
        planner_timezone=WIN_TZ_NAME
    )
    return GraphAPIClient(authenticator=authenticator, config=config)

def fetch_events_for_day(client: GraphAPIClient, target_date: datetime.date) -> List[Dict[str, Any]]:
    # Start of day in Eastern Time
    tz = ZoneInfo(TZ_NAME)
    dt_start = datetime.datetime.combine(target_date, datetime.time.min, tzinfo=tz)
    dt_end = dt_start + datetime.timedelta(days=1)
    
    # Needs to be UTC ISO strictly for Graph
    start_iso = dt_start.astimezone(datetime.timezone.utc).isoformat().replace("+00:00", "Z")
    end_iso = dt_end.astimezone(datetime.timezone.utc).isoformat().replace("+00:00", "Z")
    
    logger.info(f"Fetching calendarView for {start_iso} to {end_iso}")
    response = client.calendar_view(start_iso=start_iso, end_iso=end_iso)
    return response.get("value", [])

# --- Document Creation ---
def create_meeting_doc(event: Dict[str, Any], out_path: Path):
    doc = docx.Document()
    
    subject = event.get('subject', 'Untitled')
    start_str = event.get('start', {}).get('dateTime', '')
    end_str = event.get('end', {}).get('dateTime', '')
    organizer = event.get('organizer', {}).get('emailAddress', {}).get('name', '')
    location = event.get('location', {}).get('displayName', '')
    join_url = event.get('onlineMeeting', {}).get('joinUrl', '')
    if not join_url:
        join_url = event.get('onlineMeetingUrl', '')
        
    doc.add_heading(subject, level=1)
    
    # Metadata parsing
    sz_time = ""
    if start_str and end_str:
        try:
            s_dt = date_parser.parse(start_str)
            e_dt = date_parser.parse(end_str)
            sz_time = f"{s_dt.strftime('%Y-%m-%d %H:%M')} - {e_dt.strftime('%H:%M')}"
        except:
            sz_time = f"{start_str} - {end_str}"
            
    doc.add_paragraph(f"Meeting Time: {sz_time}")
    doc.add_paragraph(f"Organizer: {organizer}")
    if location or join_url:
        doc.add_paragraph(f"Location / Join: {location} {join_url}")
        
    doc.add_heading("Attendees", level=2)
    for att in event.get('attendees', []):
        name = att.get('emailAddress', {}).get('name', '')
        status = att.get('status', {}).get('response', 'none')
        doc.add_paragraph(f"• {name} ({status})")
        
    doc.add_heading("Description", level=2)
    body_preview = event.get('bodyPreview', '')
    if body_preview:
        doc.add_paragraph(body_preview)
        
    doc.add_paragraph(f"Event ID: {event.get('id')}").style = 'Intense Quote' # Hidden-ish
    
    doc.add_heading("Manual Notes", level=1)
    doc.add_paragraph("\n")
    
    doc.add_heading("Teams AI Summary", level=1)
    doc.add_paragraph("\n")
    
    doc.add_heading("Teams Custom Summary (Loss-Less Extraction)", level=1)
    doc.add_paragraph("\n")
    
    doc.add_heading("Otter Transcript Imports (Append-Only)", level=1)
    doc.add_paragraph("Otter transcripts will be appended below.\n")
    
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

# --- Otter Import Logic ---
def extract_text_from_file(path: Path) -> str:
    if path.suffix.lower() == '.txt':
        return path.read_text(encoding='utf-8', errors='replace')
    elif path.suffix.lower() == '.docx':
        d = docx.Document(path)
        return "\n".join([p.text for p in d.paragraphs])
    return f"[Unsupported format: {path.suffix}]"

def append_to_doc(doc_path: str, zip_name: str, transcript_text: str, comments_text: str, timestamp_str: str):
    doc = docx.Document(doc_path)
    
    # Check for existing
    search_str = f"(Source ZIP: {zip_name})"
    for p in doc.paragraphs:
        if search_str in p.text:
            logger.info(f"Transcript from {zip_name} already exists in {Path(doc_path).name}. Skipping append.")
            return

    logger.info(f"Appending transcript ({len(transcript_text)} chars) and comments ({len(comments_text)} chars) to {doc_path}")

    doc.add_heading(f"Import: {timestamp_str} (Source ZIP: {zip_name})", level=2)
    doc.add_heading("Full Transcript", level=3)
    if transcript_text.strip():
        doc.add_paragraph(transcript_text)
    else:
        doc.add_paragraph("[No transcript text found]")
    
    if comments_text.strip():
        doc.add_heading("Portions of transcript Tagged with speaker names for speaker identification", level=3)
        doc.add_paragraph(comments_text)
        
    doc.save(doc_path)
    
def get_zip_time(zip_path: Path) -> datetime.datetime:
    ts = zip_path.stat().st_mtime
    tz = ZoneInfo(TZ_NAME)
    dt = datetime.datetime.fromtimestamp(ts, tz=tz)
    return dt

def process_otter_export(zip_path: Path, index_data: Dict[str, Any], events_index: Dict[str, Any]):
    zip_name = zip_path.name
    mod_time = zip_path.stat().st_mtime
    file_size = zip_path.stat().st_size
    
    # Uniqueness key
    zip_key = f"{zip_name}_{mod_time}_{file_size}"
    if zip_key in index_data:
        logger.info(f"Skipping {zip_name} - already processed.")
        return

    logger.info(f"Processing new ZIP: {zip_name}")
    tz = ZoneInfo(TZ_NAME)
    zip_dt = get_zip_time(zip_path)
    
    transcript_text = ""
    comments_text = ""
    
    if zip_path.suffix.lower() == '.txt':
        transcript_text = extract_text_from_file(zip_path)
    else:
        with tempfile.TemporaryDirectory() as tmpdir:
            with zipfile.ZipFile(zip_path, 'r') as zf:
                zf.extractall(tmpdir)
                
            for root, _, files in os.walk(tmpdir):
                for f in files:
                    fpath = Path(root) / f
                    if f.startswith('.'): continue
                    lower_f = f.lower()
                    if 'comment' in lower_f or 'summary' in lower_f or 'takeaway' in lower_f:
                        comments_text += f"\n--- {f} ---\n" + extract_text_from_file(fpath)
                    else:
                        transcript_text += f"\n--- {f} ---\n" + extract_text_from_file(fpath)
                        
    # 1. Score candidates
    candidates = []
    for evt_id, edata in events_index.items():
        doc_path = Path(edata.get('path', ''))
        if not doc_path.exists():
            continue
            
        end_str = edata.get('end')
        if not end_str:
            continue
            
        try:
            end_dt = date_parser.parse(end_str)
            if end_dt.tzinfo is None:
                end_dt = end_dt.replace(tzinfo=tz)
                
            start_str = edata.get('start')
            if start_str:
                start_dt = date_parser.parse(start_str)
                if start_dt.tzinfo is None:
                    start_dt = start_dt.replace(tzinfo=tz)
                
                # If zip created before meeting started, impossible match
                if zip_dt < start_dt:
                    continue
        except Exception:
            continue
            
        diff = (zip_dt - end_dt).total_seconds()
        diff_hours = diff / 3600.0
        
        # Consider events that ended up to 48 hours before the export, down to -10h (spans)
        # Give highest score to events that ended just a bit before the export
        if -10 <= diff_hours <= 48:
            score = max(1, 70 - abs(diff_hours) * 1.4)  
            candidates.append({
                "score": score,
                "evt_id": evt_id, 
                "doc_path": doc_path,
                "subject": edata.get('subject', ''),
                "start": edata.get('start', ''),
                "end": edata.get('end', ''),
                "attendees": edata.get('attendees', [])
            })
            
    candidates.sort(key=lambda x: x['score'], reverse=True)
    if not candidates:
        logger.warning(f"No meeting candidate found for {zip_name} (ZIP time: {zip_dt})")
        return
        
    top_candidates = candidates[:5]
    logger.info(f"Using LLM to pick best match out of {len(top_candidates)} candidates for {zip_name}...")
    
    instructions = (
        "You are an intelligent assistant helping to disambiguate meeting transcripts. "
        "You will be given a snippet of an Otter transcript (with speaker tags if available), "
        "along with a list of recent calendar meeting candidates. "
        "Your job is to identify the ID of the meeting that best matches the transcript. "
        "Use clues like speaker names matching attendees, the subject matter, and timing. "
        "Reply ONLY with the exact Candidate ID of the matching meeting. If none seem to match well, reply with NONE."
    )
    
    candidate_list_text = "CANDIDATE MEETINGS:\n"
    for c in top_candidates:
        attn_names = []
        for a in c.get('attendees', []):
            email = a.get("emailAddress", {})
            name = email.get("name") or email.get("address", "")
            if name: attn_names.append(name)
            
        candidate_list_text += f"- ID: {c['evt_id']}\n"
        candidate_list_text += f"  Subject: {c['subject']}\n"
        candidate_list_text += f"  Time: {c['start']} to {c['end']}\n"
        candidate_list_text += f"  Attendees: {', '.join(attn_names)}\n\n"
        
    prompt = (
        f"Transcript Source file: {zip_name}\n"
        f"Export time: {zip_dt.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        f"TRANSCRIPT PREVIEW (first 4000 characters):\n{transcript_text[:4000]}\n\n"
        f"COMMENTS PREVIEW:\n{comments_text}\n\n"
        f"{candidate_list_text}"
        "Select the best matching Candidate ID. output ONLY the exact ID (no other text, no quotes)."
    )
    
    try:
        cfg = _resolve_azure_config(model_name="azure-gpt-5.2")
        client = AzureOpenAIResponsesClient(cfg)
        messages = [
            {"role": "system", "content": instructions},
            {"role": "user", "content": prompt}
        ]
        _, input_text = client.conversation_to_responses_input(messages)
        result = client.create_response(
            input_text=input_text,
            instructions=instructions,
            max_output_tokens=2000,
            reasoning_effort="low"
        )
        llm_decision = client.extract_output_text(result).strip()
        logger.info(f"LLM Raw Decision: {llm_decision}")
    except Exception as e:
        logger.warning(f"LLM matching failed: {e}. Falling back to highest time score.")
        llm_decision = top_candidates[0]["evt_id"]
        
    matched_candidate = None
    for c in top_candidates:
        if c["evt_id"] in llm_decision:
            matched_candidate = c
            break
            
    if not matched_candidate:
        logger.warning(f"LLM decided NONE or invalid ID ('{llm_decision}') for {zip_name}. Defaulting to top time proximity.")
        matched_candidate = top_candidates[0]
        
    best_score = matched_candidate["score"]
    best_evt_id = matched_candidate["evt_id"]
    target_doc_path = matched_candidate["doc_path"]
    logger.info(f"Matched {zip_name} to Doc {target_doc_path.name} (LLM Choice/Score: {best_score})")

    # 3. Append to Doc
    ts_str = zip_dt.strftime('%Y-%m-%d %H:%M:%S')
    append_to_doc(str(target_doc_path), zip_name, transcript_text.strip(), comments_text.strip(), ts_str)
    
    # 4. Mark processed
    index_data[zip_key] = {
        "zip_name": zip_name,
        "processed_at": datetime.datetime.now().isoformat(),
        "target_doc": str(target_doc_path)
    }

# --- Main Flow ---
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--date", help="Target date YYYY-MM-DD (defaults to today + yesterday)")
    parser.add_argument("--create-only", "--scaffold-only", action="store_true",
                        help="Only create meeting note docs (skip Otter processing)")
    parser.add_argument("--append-only", action="store_true",
                        help="Only process Otter exports (skip doc creation)")
    parser.add_argument("--force", action="store_true",
                        help="Regenerate docs even if they already exist (overwrites)")
    args = parser.parse_args()

    tz = ZoneInfo(TZ_NAME)
    if args.date:
        target_dates = [datetime.datetime.strptime(args.date, "%Y-%m-%d").date()]
    else:
        now_date = datetime.datetime.now(tz).date()
        target_dates = [now_date - datetime.timedelta(days=1), now_date]
        
    repo_root = Path(__file__).resolve().parent.parent

    events_index = load_index(EVENTS_INDEX_FILE)
    transcripts_index = load_index(TRANSCRIPTS_INDEX_FILE)

    if not args.append_only:
        client = get_graph_client(repo_root)
        for target_date in target_dates:
            logger.info(f"Fetching events for {target_date}...")
            events = fetch_events_for_day(client, target_date)
            logger.info(f"Found {len(events)} events for {target_date}.")
            
            day_dir = NOTES_ROOT / f"{target_date.year}" / f"{target_date.month:02d}" / f"{target_date.day:02d}"
            
            for e in events:
                eid = e.get("id")
                if not eid: continue
                
                subj = e.get("subject", "")
                subj_upper = subj.strip().upper()
                
                # Skip Canceled, Holds, and Known Non-Meeting terms
                if "Canceled:" in subj or subj_upper.startswith("HOLD") or "PTO" in subj_upper or "ENTER YOUR TIME" in subj_upper or "WATCH KIDS" in subj_upper:
                    logger.debug(f"Skipping {subj} - matches non-meeting subject pattern.")
                    continue
                    
                # Skip Out of Office / Working Elsewhere / Free
                show_as = e.get("showAs", "").lower()
                if show_as == "oof":
                    logger.debug(f"Skipping {subj} - marked as Out of Office (OOF).")
                    continue
                
                # Skip events with 0 or 1 attendee (typically just the user)
                attendees = e.get("attendees", [])
                if len(attendees) <= 1:
                    logger.debug(f"Skipping {subj} - only {len(attendees)} attendee(s).")
                    continue
                    
                eid_hash = get_event_hash(eid)
                # Duplicate check (skip if --force)
                if not args.force and eid in events_index and Path(events_index[eid]['path']).exists():
                    logger.debug(f"Skipping {e.get('subject')} - already generated.")
                    continue

                start_dt = e.get('start', {}).get('dateTime')
                end_dt = e.get('end', {}).get('dateTime')
                
                if start_dt:
                    dt_obj = date_parser.parse(start_dt)
                    if getattr(dt_obj, 'tzinfo', None):
                        dt_obj = dt_obj.astimezone(tz)
                else:
                    dt_obj = datetime.datetime.now(tz)
                    
                if not end_dt:
                    end_dt = dt_obj.isoformat()
                
                fname = build_filename(dt_obj, e.get('subject', 'Untitled'), e.get('attendees', []), eid)
                out_path = day_dir / fname
                
                if out_path.exists() and not args.force:
                    logger.info(f"Doc exists: {fname}. Skipping creation, updating index.")
                else:
                    logger.info(f"Creating doc: {fname}")
                    try:
                        create_meeting_doc(e, out_path)
                    except Exception as ex:
                        logger.error(f"Failed to create doc for event {eid_hash}: {ex}")
                        continue
                
                events_index[eid] = {
                    "hash": eid_hash,
                    "path": str(out_path),
                    "subject": e.get('subject'),
                    "start": start_dt,
                    "end": end_dt,
                    "attendees": e.get('attendees', [])
                }
            
        save_index(EVENTS_INDEX_FILE, events_index)
        
    if not args.create_only:
        if OTTER_INBOX_DIR.exists():
            for ext in ("*.zip", "*.txt"):
                for zpath in OTTER_INBOX_DIR.glob(ext):
                    try:
                        process_otter_export(zpath, transcripts_index, events_index)
                    except Exception as e:
                        logger.error(f"Failed to process {zpath.name}: {e}")
            save_index(TRANSCRIPTS_INDEX_FILE, transcripts_index)
        else:
            logger.warning(f"Otter inbox not found at {OTTER_INBOX_DIR}")

if __name__ == "__main__":
    main()
