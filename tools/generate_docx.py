#!/usr/bin/env python3
"""
Generate DOCX from Markdown using a .dotx Template
--------------------------------------------------
This utility converts a Markdown file into a Word document (.docx) using a specified
Word Template (.dotx). Ideally suited for generating Statements of Work (SOWs) 
or other formal documents from Markdown drafts.

Key Features:
- Handles .dotx templates (which python-docx doesn't natively support) by temporary conversion.
- Clears default content from the template (Latin placeholder text) while preserving 
  page properties (margins, orientation, size).
- Maps Markdown headers (#, ##) to Word Styles (Heading 1, Heading 2).
- NOTE: Maps H1 (#) to 'Heading 1' instead of 'Title' to avoid missing style errors.

Usage:
    python tools/generate_docx.py --input <input.md> --template <template.dotx> --output <output.docx>

Dependencies:
    pip install python-docx
"""

import re
import os
import shutil
import zipfile
import tempfile
import argparse
import sys
from docx import Document

def convert_dotx_to_docx(dotx_path):
    """
    Converts a .dotx template to a temporary .docx file.
    Required because python-docx cannot open .dotx files directly.
    """
    # Create a temp dir for extraction
    extract_dir = tempfile.mkdtemp()
    
    # Create a temp file path for the output docx (outside the extract dir)
    temp_docx_handle, temp_docx_path = tempfile.mkstemp(suffix='.docx')
    os.close(temp_docx_handle)
    
    # Extract dotx
    print(f"Extracting template from {dotx_path}...")
    with zipfile.ZipFile(dotx_path, 'r') as zin:
        zin.extractall(extract_dir)
        
    # Modify [Content_Types].xml to change MIME type from template to document
    ct_path = os.path.join(extract_dir, "[Content_Types].xml")
    if os.path.exists(ct_path):
        with open(ct_path, 'r') as f:
            content = f.read()
        
        content = content.replace(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
        )
        
        with open(ct_path, 'w') as f:
            f.write(content)
    else:
        print("Warning: [Content_Types].xml not found in template extraction")
        
    # Zip back to docx
    with zipfile.ZipFile(temp_docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                full_path = os.path.join(root, file)
                rel_path = os.path.relpath(full_path, extract_dir)
                zout.write(full_path, rel_path)
    
    # Clean up extraction dir
    shutil.rmtree(extract_dir)
                
    return temp_docx_path

def add_markdown_paragraph(doc, text, style='Normal'):
    """Adds a paragraph with basic bold markdown support (**text**)."""
    try:
        p = doc.add_paragraph(style=style)
    except Exception as e:
        print(f"Warning: Style '{style}' not found. Falling back to default. Error: {e}")
        p = doc.add_paragraph() 
        
    # Simple bold parser: "Some **bold** text"
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def markdown_to_docx(md_path, template_path, output_path):
    """Main conversion logic."""
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown file not found: {md_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    # Convert template
    working_template = convert_dotx_to_docx(template_path)
    print(f"Converted temporary template: {working_template}")
    
    doc = Document(working_template)

    # Clear existing content from template (e.g. placeholder text)
    # process body elements in reverse order to avoid index issues
    element = doc._body._element
    if element is not None:
        for child in list(element):
             # Preserve section properties (sectPr) to keep margins/page setup
             if child.tag.endswith('sectPr'):
                 continue
             element.remove(child)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_mode = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Table Start / Continue
        if line.startswith('|'):
            table_mode = True
            # Parse row: remove outer pipes and split
            row_content = [c.strip() for c in line.strip('|').split('|')]
            table_data.append(row_content)
            continue
            
        # Table End (if we were in table mode but line doesn't start with |)
        elif table_mode:
            table_mode = False
            # Create table in doc
            # Filter out separator lines (e.g. ---|---)
            data_rows = [row for row in table_data if not set(row[0]) <=set('- :')]
            
            if data_rows:
                # Create table
                cols = len(data_rows[0])
                rows = len(data_rows)
                
                # Check for styles - usually "Table Grid" or "Grid Table" exists
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for r_idx, row in enumerate(data_rows):
                    # Ensure we don't exceed columns if row has extra pipes
                    safe_cols = min(len(row), cols)
                    for c_idx in range(safe_cols):
                        cell_text = row[c_idx]
                        # Handle <br> in tables
                        cell_text = cell_text.replace("<br>", "\n")
                        cell = table.cell(r_idx, c_idx)
                        # Add text with formatting support (remove default empty para)
                        p = cell.paragraphs[0]
                        p.text = "" 
                        # Bold logic for cells
                        parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
            table_data = []
            doc.add_paragraph() # Spacer after table

        if not line:
            continue

        # Headers
        if line.startswith('# '):
             # Title - mapped directly to Heading 1 to avoid missing 'Title' style
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Bullet Points
        elif line.startswith('* ') or line.startswith('- '):
            add_markdown_paragraph(doc, line[2:], style='List Bullet')
        
        # Blockquotes
        elif line.startswith('> '):
            add_markdown_paragraph(doc, line[2:], style='Quote')
            
        # Normal Text
        else:
            add_markdown_paragraph(doc, line, style='Normal')

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f"Success! Document generated at: {output_path}")
    
    # Cleanup temp template
    if os.path.exists(working_template):
        os.remove(working_template)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Markdown to Docx using a Dotx template.")
    parser.add_argument("--input", required=True, help="Path to input Markdown file")
    parser.add_argument("--template", required=True, help="Path to .dotx template file")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    
    args = parser.parse_args()
    
    try:
        markdown_to_docx(args.input, args.template, args.output)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
