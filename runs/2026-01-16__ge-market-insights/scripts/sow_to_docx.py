import re
import os
import shutil
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt
import sys

# Paths
RUN_DIR = "/Users/matt_bossemeyer/Projects/wt-2026-01-16__ge-market-insights/runs/2026-01-16__ge-market-insights"
INPUT_MD = os.path.join(RUN_DIR, "exports/draft_sow_ge_pilot_v8.md")
TEMPLATE_DOTX = os.path.join(RUN_DIR, "inputs/26-Legal-WordTemp.dotx")
OUTPUT_DOCX = os.path.join(RUN_DIR, "exports/GE_Market_Insights_Pilot_SOW_v8.docx")

def convert_dotx_to_docx(dotx_path):
    # Create a temp dir for extraction
    extract_dir = tempfile.mkdtemp()
    
    # Create a temp file path for the output docx (outside the extract dir)
    temp_docx_handle, temp_docx_path = tempfile.mkstemp(suffix='.docx')
    os.close(temp_docx_handle)
    
    # Extract dotx
    print(f"Extracting template from {dotx_path} to {extract_dir}...")
    with zipfile.ZipFile(dotx_path, 'r') as zin:
        zin.extractall(extract_dir)
        
    # Modify [Content_Types].xml
    ct_path = os.path.join(extract_dir, "[Content_Types].xml")
    if os.path.exists(ct_path):
        with open(ct_path, 'r') as f:
            content = f.read()
        
        content = content.replace(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
        )
        
        with open(ct_path, 'w') as f:
            f.write(content)
    else:
        print("Warning: [Content_Types].xml not found in template extraction")
        
    # Zip back to docx
    print(f"Repackaging to {temp_docx_path}...")
    with zipfile.ZipFile(temp_docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                full_path = os.path.join(root, file)
                rel_path = os.path.relpath(full_path, extract_dir)
                zout.write(full_path, rel_path)
    
    # Clean up extraction dir
    shutil.rmtree(extract_dir)
                
    return temp_docx_path

def add_markdown_paragraph(doc, text, style='Normal'):
    try:
        p = doc.add_paragraph(style=style)
    except:
        p = doc.add_paragraph() # Fallback if style missing
        
    # Simple bold parser: "Some **bold** text"
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def markdown_to_docx(md_path, template_path, output_path):
    # Convert template
    working_template = convert_dotx_to_docx(template_path)
    print(f"Converted template to: {working_template}")
    
    doc = Document(working_template)

    # Clear existing content from template (e.g. placeholder text)
    # process body elements in reverse order to avoid index issues
    element = doc._body._element
    if element is not None:
        for child in list(element):
             # Preserve section properties (sectPr) to keep margins/page setup
             if child.tag.endswith('sectPr'):
                 continue
             element.remove(child)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_mode = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Table handling
        if line.startswith('|'):
            table_mode = True
            # Parse row: remove outer pipes and split
            row_content = [c.strip() for c in line.strip('|').split('|')]
            table_data.append(row_content)
            continue
        elif table_mode:
            # Table ended
            table_mode = False
            # Create table in doc
            # Filter out separator lines (e.g. ---|---)
            data_rows = [row for row in table_data if not set(row[0]) <=set('- :')]
            
            if data_rows:
                table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]))
                table.style = 'Table Grid'
                for r_idx, row in enumerate(data_rows):
                    for c_idx, cell_text in enumerate(row):
                        # Handle <br> in tables
                        cell_text = cell_text.replace("<br>", "\n")
                        cell = table.cell(r_idx, c_idx)
                        # Add text with formatting support
                        p = cell.paragraphs[0]
                        # Bold logic for cells
                        parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
            table_data = []
            doc.add_paragraph() # Spacer

        if not line:
            continue

        # Headers
        if line.startswith('# '):
             # Title - mapped directly to Heading 1 to avoid missing 'Title' style
             # (original logic tried level=0 which requires 'Title')
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Bullet Points
        elif line.startswith('* ') or line.startswith('- '):
            add_markdown_paragraph(doc, line[2:], style='List Bullet')
        
        # Blockquotes?
        elif line.startswith('> '):
            add_markdown_paragraph(doc, line[2:], style='Quote')
            
        # Normal Text
        else:
            add_markdown_paragraph(doc, line, style='Normal')

    doc.save(output_path)
    print(f"Generated SOW at: {output_path}")

if __name__ == "__main__":
    print("Starting script...")
    try:
        markdown_to_docx(INPUT_MD, TEMPLATE_DOTX, OUTPUT_DOCX)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
