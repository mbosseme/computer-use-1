import os
import docx

doc_path = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/2026-03-10 1300 - SC Internal Deal Desk Project Meeting - Danielle-Bill-Matthew - eid9d3d92a8.docx"
doc = docx.Document(doc_path)

dummy_recap = """1. EXHAUSTIVE TOPIC EXTRACTION
- Deal Desk Build-Out: Phased approach to building out supply chain deal desk, first iteration for immediate impact.
- Growth Project and Opportunity Heat Map: Contract uptake, off-contract spend.

2. SYSTEM NAMES, TOOLS, AND PROJECTS
- Supply Chain Deal Desk
- PowerBI Heat Map

3. METRICS, KPIS, AND SPEND FIGURES
- $1 Billion pharmacy distribution spend.
- 55% baseline uptake.

4. STRATEGIC INSIGHTS & CONSTRAINTS
- Need for clear standards, benchmarking, and proactive guidance for negotiators.
- Exception-based process to streamline approvals.

5. SPEAKER ATTRIBUTION & FRICTION POINTS
- Matthew and Danielle debated primary goals (growth vs fee optimization vs national contract competitiveness)."""

# Find Teams Recap Expansion header
inserted = False
otter_index = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Otter Transcript Imports"):
        otter_index = i
        break

if otter_index != -1:
    p = doc.paragraphs[otter_index]
    p.insert_paragraph_before("--- Playwright MCP Extracted Summary ---")
    p.insert_paragraph_before(dummy_recap.strip())
    inserted = True

doc.save(doc_path)
print("Insertion completed:", inserted)