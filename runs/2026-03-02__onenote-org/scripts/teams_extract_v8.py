import sys
import os
import time
from playwright.sync_api import sync_playwright
import docx

def format_transcript(raw_text):
    """Clean up the raw Teams body dump to look like a usable transcript."""
    # This strips out obvious top-level Teams nav junk
    bad_lines = {"Activity", "Chat", "Teams", "Calendar", "Calls", "OneDrive", "Mentions", "Transcript"}
    lines = raw_text.split('\n')
    cleaned = []
    
    # Simple heuristic to find where the actual meeting data starts.
    # Often it includes standard Teams UI we want to skip.
    capture = False
    for line in lines:
        t = line.strip()
        if not t: continue
        
        # Once we see a timecode (like 00:00) or 'Transcript', we know we're in the right area
        if "Transcript" in t and "AI summary" in raw_text:
            capture = True
            
        if t in bad_lines: continue
        
        cleaned.append(t)
    return "\n".join(cleaned)

def main():
    profile_dir = os.path.abspath('runs/2026-03-02__onenote-org/playwright-profile')
    
    summary_text = ""
    transcript_text = ""

    print("\n" + "="*80)
    print("STARTING TEAMS AUTO-EXTRACTOR")
    print("="*80)
    
    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=profile_dir,
            headless=False,
            viewport={"width": 1440, "height": 900}
        )
        
        # Connect to existing page or make a new one
        page = browser.pages[0] if browser.pages else browser.new_page()
        
        # Note: If you're already on the page, this just refreshes or ensures we are in Teams.
        if "teams.microsoft.com" not in page.url:
            print("Navigating to Teams...")
            page.goto("https://teams.microsoft.com/v2/", wait_until="domcontentloaded")
        
        print("\n⏳ ACTION 1: AI SUMMARY DATA")
        print("-> Go to the 'SC Internal Deal Desk Project Meeting' recording page.")
        print("-> Make sure the 'AI summary' tab is selected.")
        print("-> Click 'Expand all' under the Meeting notes if available.")
        print("\nWaiting 20 seconds for you to get ready...")
        
        for i in range(20, 0, -1):
            sys.stdout.write(f"\rTime remaining: {i} seconds... ")
            sys.stdout.flush()
            time.sleep(1)
            
        print("\n\n📸 Capturing AI Summary...")
        try:
            summary_text = page.locator("body").inner_text()
            print("   ✅ Summary captured!")
        except Exception as e:
            print(f"   ❌ Error grabbing summary: {e}")

        print("\n⏳ ACTION 2: TRANSCRIPT DATA")
        print("-> Click the 'Transcript' tab now.")
        print("-> Scroll down a tiny bit to make sure it loads.")
        print("\nWaiting 15 seconds...")
        for i in range(15, 0, -1):
            sys.stdout.write(f"\rTime remaining: {i} seconds... ")
            sys.stdout.flush()
            time.sleep(1)

        print("\n\n📸 Capturing Transcript...")
        try:
            transcript_text = page.locator("body").inner_text()
            print("   ✅ Transcript captured!")
        except Exception as e:
            print(f"   ❌ Error grabbing transcript: {e}")

        # Quick screenshot for verification
        os.makedirs("runs/2026-03-02__onenote-org/tmp", exist_ok=True)
        try:
            page.screenshot(path="runs/2026-03-02__onenote-org/tmp/teams_last_state.png")
        except:
            pass

        print("Closing browser...")
        browser.close()

    print("\nProcessing captured text...")
    
    # Basic isolation logic 
    clean_summary = ""
    try:
        # Split out the left-rail navigation noise if possible
        if "Meeting notes" in summary_text:
            parts = summary_text.split("Meeting notes")
            clean_summary = "Meeting notes" + parts[-1]
            if "Transcript" in clean_summary:
                clean_summary = clean_summary.split("Transcript")[0]
        else:
            clean_summary = summary_text
    except:
        clean_summary = summary_text

    clean_transcript = format_transcript(transcript_text)

    # Output to TMP for debugging
    with open("runs/2026-03-02__onenote-org/tmp/teams_ai_backup.txt", "w", encoding="utf-8") as f:
        f.write(summary_text)
    with open("runs/2026-03-02__onenote-org/tmp/teams_transcript_backup.txt", "w", encoding="utf-8") as f:
        f.write(transcript_text)

    # ---------------------------------------------------------
    # FIND AND UPDATE THE WORD DOCUMENT
    # ---------------------------------------------------------
    print("\nLocating the target Word document...")
    docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
    target_doc_path = None
    if os.path.exists(docs_dir):
        for file in os.listdir(docs_dir):
            if "SC Internal Deal Desk Project" in file and file.endswith(".docx") and not file.startswith("~$"):
                target_doc_path = os.path.join(docs_dir, file)
                break
    
    if not target_doc_path:
        print(f"❌ Could not find target doc in {docs_dir}")
        return
        
    print(f"Updating: {os.path.basename(target_doc_path)}")
    try:
        document = docx.Document(target_doc_path)
        
        # Idempotency check
        already_added = any("Teams AI Summary & Transcript" in p.text for p in document.paragraphs)
        if already_added:
            print("Looks like this Teams data was already added previously. Skipping append.")
            return
            
        # Append data to the end of the Word Document
        document.add_heading("Teams AI Summary & Transcript Extract", level=1)
        
        document.add_heading("AI Summary Notes", level=2)
        if len(clean_summary) > 50:
            for line in clean_summary.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Failed to cleanly extract AI summary text]")
            
        document.add_heading("Full Raw Transcript", level=2)
        if len(clean_transcript) > 50:
            for chunk in clean_transcript.split('\n'):
                if chunk.strip():
                    document.add_paragraph(chunk.strip())
        else:
            document.add_paragraph("[Failed to cleanly extract transcript text]")
            
        document.save(target_doc_path)
        print("\n🎉 SUCCESS! I have saved the Team's AI Summary and Transcript into your meeting notes doc.")
        
    except Exception as e:
        print(f"❌ Error handling file update: {e}")

if __name__ == "__main__":
    main()
