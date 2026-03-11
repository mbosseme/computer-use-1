import sys
import os
import time
from playwright.sync_api import sync_playwright
import docx

def main():
    profile_dir = os.path.abspath('runs/2026-03-02__onenote-org/playwright-profile')
    
    summary_text = ""
    transcript_text = ""

    print("Starting Playwright Automator...")
    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=profile_dir,
            headless=False,
            viewport={"width": 1440, "height": 900}
        )
        # Connect to existing page
        page = browser.pages[0] if browser.pages else browser.new_page()
        
        print("I am keeping the browser open for 15 seconds. Please ensure you are on the SC Internal Deal Desk recording page, on the 'AI summary' tab.")
        time.sleep(15)
        
        print("Attempting Extraction...")
        
        try:
            print("Extracting AI summary text...")
            # Ensure "Expand all" is clicked if visible
            expand_btn = page.get_by_role("button", name="Expand all").first
            if expand_btn.is_visible():
                expand_btn.click()
                time.sleep(1)

            # Try to narrow to the meeting notes right pane
            panel = page.locator('div[role="tabpanel"]').first
            if panel.is_visible():
                summary_text = panel.inner_text()
            else:
                summary_text = page.locator("body").inner_text()
        except Exception as e:
            print("Error during AI summary extraction:", e)

        print("Looking for the 'Transcript' tab and clicking it...")
        try:
            transcript_tab = page.get_by_role("tab", name="Transcript").first
            if not transcript_tab.is_visible():
                transcript_tab = page.locator('text="Transcript"').first
                
            if transcript_tab.is_visible():
                transcript_tab.click()
                print("Clicked Transcript tab. Waiting 5 seconds for load...")
                time.sleep(5) 
                
                print("Extracting Transcript text...")
                t_panel = page.locator('div[role="tabpanel"]').first
                if t_panel.is_visible():
                    transcript_text = t_panel.inner_text()
                else:
                    transcript_text = page.locator("body").inner_text()
            else:
                print("Could not click Transcript tab.")
                
        except Exception as e:
             print("Error during Transcript extraction:", e)

        # Grab a screenshot for debugging
        os.makedirs("runs/2026-03-02__onenote-org/tmp", exist_ok=True)
        try:
            page.screenshot(path="runs/2026-03-02__onenote-org/tmp/teams_last_state.png")
        except:
            pass

        browser.close()

    print(f"Extraction complete. Summary size: {len(summary_text)} | Transcript size: {len(transcript_text)}")
    
    with open("runs/2026-03-02__onenote-org/tmp/teams_ai_backup.txt", "w") as f:
        f.write(summary_text)
    with open("runs/2026-03-02__onenote-org/tmp/teams_transcript_backup.txt", "w") as f:
        f.write(transcript_text)

    if len(summary_text) < 10 and len(transcript_text) < 10:
        print("Failed to extract meaningful text. Exiting.")
        return

    print("\nLocating matching Word document...")
    docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
    target_doc_path = None
    if os.path.exists(docs_dir):
        for file in os.listdir(docs_dir):
            if "SC Internal Deal Desk Project" in file and file.endswith(".docx") and not file.startswith("~$"):
                target_doc_path = os.path.join(docs_dir, file)
                break
    
    if not target_doc_path:
        print(f"Could not find target Word document in {docs_dir}")
        return
        
    print(f"Updating doc: {os.path.basename(target_doc_path)}")
    try:
        document = docx.Document(target_doc_path)
        
        already_added = any("Teams AI Summary & Transcript" in p.text for p in document.paragraphs)
        if already_added:
            print("Content already present in document. Aborting to avoid duplicates.")
            return
            
        document.add_heading("Teams AI Summary & Transcript Imports", level=1)
        
        document.add_heading("Teams AI Summary", level=2)
        if summary_text:
            for line in summary_text.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Failed to extract AI summary]")
            
        document.add_heading("Full Teams Transcript", level=2)
        if transcript_text:
            # Word limits paragraph size, so split by lines
            for line in transcript_text.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Failed to extract transcript]")
            
        document.save(target_doc_path)
        print("Document successfully updated and saved to OneDrive!")
        
    except Exception as e:
        print(f"Error handling docx: {e}")

if __name__ == "__main__":
    main()
