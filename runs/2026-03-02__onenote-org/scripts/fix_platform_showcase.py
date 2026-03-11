"""
Fix up the Platform Showcase doc: move AI summary content under
'Screenshots / Images', and reorganize Teams Recap content properly.
"""
from docx import Document
from pathlib import Path

doc_path = '/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/2026-03-10 0930 - Platform Showcase - Zach-Elise-Jesse - eid1f0145bc.docx'
ai_file = 'runs/2026-03-02__onenote-org/tmp/platform_showcase_ai_summary.txt'
cs_file = 'runs/2026-03-02__onenote-org/tmp/platform_showcase_custom_summary.txt'

doc = Document(doc_path)

# Collect text content first
ai_text = Path(ai_file).read_text(encoding='utf-8')
cs_text = Path(cs_file).read_text(encoding='utf-8')
ai_lines = [l for l in ai_text.split('\n') if l.strip()]
cs_lines = [l for l in cs_text.split('\n') if l.strip()]

# Now rebuild: remove the misplaced AI Summary & Screenshots heading and
# the Teams Recap that was appended at end. We'll re-inject properly.

# Find paragraphs to remove (index 36 onwards to end = the fallback appended content)
# Keep paragraphs 0-35 (through Otter Transcript heading), remove 36+
paras_to_remove = []
for i in range(len(doc.paragraphs) - 1, 35, -1):
    paras_to_remove.append(i)

# Remove from end to avoid index shifting
for i in sorted(paras_to_remove, reverse=True):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# Now inject AI summary after "Screenshots / Images" (currently index 33)
# After removal, index 34 is the empty Normal paragraph, index 35 is Otter H1
# We want to insert before index 35 (Otter H1)
otter_heading = doc.paragraphs[35]._element
for line in reversed(ai_lines):
    new_p = doc.add_paragraph(line)
    otter_heading.addprevious(new_p._element)

# Now add Teams Recap Extraction section at the very end
doc.add_heading("Teams Recap Extraction (Append-Only)", level=1)
for line in cs_lines:
    doc.add_paragraph(line)

doc.save(doc_path)

# Verify
doc2 = Document(doc_path)
print("Final structure:")
for i, p in enumerate(doc2.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f"  [{i}] {p.style.name}: {p.text[:100]}")
print(f"Total: {len(doc2.paragraphs)} paragraphs")
