"""
Inject AI summary with embedded images from rich clipboard HTML into a meeting notes .docx.

Usage:
    python inject_ai_summary_rich.py <doc_path> [--from-clipboard | --from-file <html_file>]

The script reads HTML from the macOS clipboard (via osascript) or from a saved HTML file,
parses text and embedded base64 images, and injects them under the
"AI Summary & Screenshots" or "Screenshots / Images" heading in the Word doc.
"""
import argparse
import base64
import io
import re
import subprocess
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches


def get_html_from_clipboard():
    """Read HTML clipboard content via macOS osascript."""
    r = subprocess.run(
        ['osascript', '-e', 'the clipboard as «class HTML»'],
        capture_output=True,
    )
    if r.returncode != 0:
        print(f"Error reading clipboard: {r.stderr.decode()}", file=sys.stderr)
        return None

    raw = r.stdout.decode('utf-8', errors='replace').strip()

    # osascript returns: «data HTML[hex]»
    if raw.startswith('«data HTML') and raw.endswith('»'):
        hex_str = raw[10:-1]
    else:
        print("Unexpected clipboard format", file=sys.stderr)
        return None

    try:
        return bytes.fromhex(hex_str).decode('utf-8', errors='replace')
    except ValueError:
        print("Failed to decode hex", file=sys.stderr)
        return None


def extract_content_blocks(html: str):
    """
    Parse HTML from Teams AI summary into ordered blocks of text and images.

    Returns a list of dicts:
      {"type": "text", "text": "...", "bold": bool}
      {"type": "image", "data": bytes, "ext": "png"|"jpeg", "alt": "..."}
    """
    soup = BeautifulSoup(html, 'html.parser')
    blocks = []

    def walk(element, depth=0):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    # Check if parent is bold
                    is_bold = child.parent and child.parent.name == 'b'
                    blocks.append({"type": "text", "text": text, "bold": is_bold, "depth": depth})
            elif isinstance(child, Tag):
                if child.name == 'img':
                    src = child.get('src', '')
                    alt = child.get('alt', '')
                    m = re.match(r'data:image/([\w+]+);base64,(.+)', src, re.DOTALL)
                    if m:
                        ext = m.group(1).replace('jpeg', 'jpg')
                        img_bytes = base64.b64decode(m.group(2))
                        blocks.append({
                            "type": "image",
                            "data": img_bytes,
                            "ext": ext,
                            "alt": alt,
                        })
                elif child.name == 'br':
                    pass  # skip line breaks
                elif child.name in ('ul', 'ol'):
                    walk(child, depth + 1)
                elif child.name == 'li':
                    walk(child, depth)
                elif child.name in ('div', 'p', 'span', 'b', 'i', 'em', 'strong'):
                    walk(child, depth)
                else:
                    walk(child, depth)

    walk(soup)
    return blocks


def merge_text_blocks(blocks):
    """
    Merge consecutive text blocks at the same depth into single paragraphs.
    Images stay as separate blocks.
    """
    merged = []
    for b in blocks:
        if b["type"] == "text":
            # If the previous block is text at the same depth, merge
            if (merged
                    and merged[-1]["type"] == "text"
                    and merged[-1]["depth"] == b["depth"]
                    and not merged[-1].get("_is_header", False)):
                # Append as a text run
                merged[-1].setdefault("runs", [])
                if not merged[-1]["runs"]:
                    # Convert existing to first run
                    merged[-1]["runs"] = [{"text": merged[-1]["text"], "bold": merged[-1]["bold"]}]
                merged[-1]["runs"].append({"text": b["text"], "bold": b["bold"]})
                merged[-1]["text"] += " " + b["text"]
            else:
                merged.append(dict(b))
        else:
            merged.append(b)
    return merged


def find_target_heading(doc):
    """Find the paragraph index of the AI Summary heading."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            if any(term in p.text for term in [
                "Teams AI Summary",
                "AI Summary & Screenshots",
                "Screenshots / Images",
            ]):
                return i
    return None


def find_next_heading(doc, target_idx):
    """Find the next heading at same or higher level after target."""
    target_style = doc.paragraphs[target_idx].style.name
    parts = target_style.split()
    target_level = int(parts[-1]) if parts[-1].isdigit() else 1

    for i in range(target_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading'):
            parts = p.style.name.split()
            lvl = int(parts[-1]) if parts[-1].isdigit() else 99
            if lvl <= target_level:
                return i
    return None


def clear_section(doc, start_idx, end_idx):
    """Remove existing content between start heading and next heading."""
    # Remove paragraphs between start+1 and end (exclusive)
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    elements_to_remove = []
    for i in range(start_idx + 1, end_idx):
        elements_to_remove.append(doc.paragraphs[i]._element)

    for el in elements_to_remove:
        el.getparent().remove(el)


def inject_blocks(doc, doc_path, blocks, replace=True):
    """Inject content blocks into the doc under the AI summary heading."""
    target_idx = find_target_heading(doc)

    if target_idx is None:
        print("WARNING: No AI Summary heading found, creating one")
        doc.add_heading("Teams AI Summary", level=1)
        target_idx = len(doc.paragraphs) - 1

    next_heading_idx = find_next_heading(doc, target_idx)

    # Optionally clear existing content in this section
    if replace:
        clear_section(doc, target_idx, next_heading_idx)
        # After clearing, the next heading (if any) is now at target_idx + 1
        insert_before_idx = target_idx + 1 if next_heading_idx is not None else None
    else:
        insert_before_idx = next_heading_idx

    # Insert blocks
    # We need to insert using lxml element manipulation for correct positioning.
    # addprevious() inserts immediately before the reference element, so iterating
    # in forward order preserves the original block sequence.
    if insert_before_idx is not None and insert_before_idx < len(doc.paragraphs):
        ref_element = doc.paragraphs[insert_before_idx]._element
        for block in blocks:
            if block["type"] == "text":
                p = doc.add_paragraph()
                p.clear()
                if block.get("runs"):
                    for run_info in block["runs"]:
                        run = p.add_run(run_info["text"] + " ")
                        run.bold = run_info["bold"]
                else:
                    run = p.add_run(block["text"])
                    run.bold = block.get("bold", False)
                ref_element.addprevious(p._element)
            elif block["type"] == "image":
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(io.BytesIO(block["data"]), width=Inches(5.5))
                ref_element.addprevious(p._element)
    else:
        # Append at end
        for block in blocks:
            if block["type"] == "text":
                p = doc.add_paragraph()
                p.clear()
                if block.get("runs"):
                    for run_info in block["runs"]:
                        run = p.add_run(run_info["text"] + " ")
                        run.bold = run_info["bold"]
                else:
                    run = p.add_run(block["text"])
                    run.bold = block.get("bold", False)
            elif block["type"] == "image":
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(io.BytesIO(block["data"]), width=Inches(5.5))

    doc.save(doc_path)

    text_count = sum(1 for b in blocks if b["type"] == "text")
    img_count = sum(1 for b in blocks if b["type"] == "image")
    print(f"SAVED: Injected {text_count} text blocks + {img_count} images into {Path(doc_path).name}")


def main():
    parser = argparse.ArgumentParser(description="Inject rich AI summary into meeting notes docx")
    parser.add_argument("doc_path", help="Path to the .docx file")
    parser.add_argument("--from-clipboard", action="store_true", help="Read from macOS clipboard")
    parser.add_argument("--from-file", type=str, help="Read from saved HTML file")
    parser.add_argument("--no-replace", action="store_true",
                        help="Append instead of replacing existing AI summary content")
    parser.add_argument("--save-html", type=str,
                        help="Save the raw HTML to this path (for debugging/re-use)")
    args = parser.parse_args()

    # Get HTML
    if args.from_file:
        html = Path(args.from_file).read_text(encoding='utf-8')
    elif args.from_clipboard:
        html = get_html_from_clipboard()
        if not html:
            print("Failed to read clipboard", file=sys.stderr)
            sys.exit(1)
    else:
        # Default to clipboard
        html = get_html_from_clipboard()
        if not html:
            print("Failed to read clipboard", file=sys.stderr)
            sys.exit(1)

    # Optionally save HTML for re-use
    if args.save_html:
        Path(args.save_html).parent.mkdir(parents=True, exist_ok=True)
        Path(args.save_html).write_text(html, encoding='utf-8')
        print(f"Saved raw HTML to {args.save_html}")

    # Parse
    blocks = extract_content_blocks(html)
    blocks = merge_text_blocks(blocks)

    text_count = sum(1 for b in blocks if b["type"] == "text")
    img_count = sum(1 for b in blocks if b["type"] == "image")
    print(f"Parsed: {text_count} text blocks, {img_count} images")

    if text_count == 0 and img_count == 0:
        print("No content found in HTML", file=sys.stderr)
        sys.exit(1)

    # Inject into doc
    doc = Document(args.doc_path)
    inject_blocks(doc, args.doc_path, blocks, replace=not args.no_replace)


if __name__ == "__main__":
    main()
