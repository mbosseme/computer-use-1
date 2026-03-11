"""
Inject Teams Recap data (AI Summary + Custom Summary) into a Word document.

Usage:
  python3 runs/2026-03-02__onenote-org/scripts/inject_teams_recap.py \
    --doc "path/to/meeting.docx" \
    --ai-summary "path/to/ai_summary.txt" \
    --custom-summary "path/to/custom_summary.txt"
"""
import argparse
import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt


def find_heading_index(doc, heading_text, level=None):
    """Find the paragraph index of a heading by text match."""
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            if heading_text.lower() in para.text.lower():
                if level is None or para.style.name == f'Heading {level}':
                    return i
    return None


def has_existing_recap(doc, marker_prefix="Teams Recap Import:"):
    """Check if recap data was already injected (idempotency)."""
    for para in doc.paragraphs:
        if para.text.startswith(marker_prefix):
            return True
    return False


def inject_after_heading(doc, heading_text, level, content_lines, sub_heading=None):
    """Insert content paragraphs after a specific heading."""
    idx = find_heading_index(doc, heading_text, level)
    if idx is None:
        return False

    insert_at = idx + 1
    # Skip any existing content until next heading of same or higher level
    # We want to append at the end of this section
    while insert_at < len(doc.paragraphs):
        p = doc.paragraphs[insert_at]
        if p.style.name.startswith('Heading'):
            style_level = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 99
            if style_level <= level:
                break
        insert_at += 1

    # Insert content at the found position
    ref_para = doc.paragraphs[insert_at - 1] if insert_at > 0 else doc.paragraphs[0]

    if sub_heading:
        new_heading = ref_para._element.addnext(
            doc.add_heading(sub_heading, level=level + 1)._element
        )

    for line in reversed(content_lines):
        if line.strip():
            new_para = ref_para._element.addnext(
                doc.add_paragraph(line.strip())._element
            )

    return True


def main():
    parser = argparse.ArgumentParser(description="Inject Teams recap into docx")
    parser.add_argument("--doc", required=True, help="Path to the Word document")
    parser.add_argument("--ai-summary", help="Path to AI summary text file")
    parser.add_argument("--custom-summary", help="Path to custom summary text file")
    args = parser.parse_args()

    doc_path = Path(args.doc)
    if not doc_path.exists():
        print(f"ERROR: Document not found: {doc_path}")
        return

    doc = Document(str(doc_path))
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    marker = f"Teams Recap Import: {timestamp}"
    changes = False

    # --- Inject AI Summary under "Teams AI Summary" H1 (or legacy names) ---
    if args.ai_summary:
        ai_path = Path(args.ai_summary)
        if ai_path.exists():
            ai_text = ai_path.read_text(encoding="utf-8")
            # Try new name first, then legacy names
            idx = find_heading_index(doc, "Teams AI Summary")
            if idx is None:
                idx = find_heading_index(doc, "AI Summary & Screenshots")
            if idx is None:
                idx = find_heading_index(doc, "Screenshots / Images")
            target_level = int(doc.paragraphs[idx].style.name.split()[-1]) if idx is not None else 1

            if idx is not None:
                # Find the next heading at same or higher level
                insert_before = None
                for i in range(idx + 1, len(doc.paragraphs)):
                    p = doc.paragraphs[i]
                    if p.style.name.startswith('Heading'):
                        style_level = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 99
                        if style_level <= target_level:
                            insert_before = i
                            break

                lines_to_add = [l for l in ai_text.split('\n') if l.strip()]
                ref_element = doc.paragraphs[insert_before]._element if insert_before else None

                if ref_element is not None:
                    for line in lines_to_add:
                        new_p = doc.add_paragraph(line.strip())
                        ref_element.addprevious(new_p._element)
                else:
                    for line in lines_to_add:
                        doc.add_paragraph(line.strip())

                print(f"  Injected AI Summary ({len(lines_to_add)} lines)")
                changes = True
            else:
                print("  WARNING: AI Summary heading not found, creating one")
                doc.add_heading("Teams AI Summary", level=1)
                for line in ai_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                changes = True

    # --- Inject Custom Summary under "Teams Custom Summary" H1 (or legacy name) ---
    if args.custom_summary:
        cs_path = Path(args.custom_summary)
        if cs_path.exists():
            cs_text = cs_path.read_text(encoding="utf-8")
            # Try new name first, then legacy
            idx = find_heading_index(doc, "Teams Custom Summary")
            if idx is None:
                idx = find_heading_index(doc, "Teams Recap Extraction")
            if idx is not None:
                # Find the next H1 heading after this section
                insert_before = None
                for i in range(idx + 1, len(doc.paragraphs)):
                    p = doc.paragraphs[i]
                    if p.style.name.startswith('Heading'):
                        style_level = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 99
                        if style_level <= 1:
                            insert_before = i
                            break

                lines_to_add = [l for l in cs_text.split('\n') if l.strip()]
                ref_element = doc.paragraphs[insert_before]._element if insert_before else None

                if ref_element is not None:
                    for line in lines_to_add:
                        new_p = doc.add_paragraph(line.strip())
                        ref_element.addprevious(new_p._element)
                else:
                    for line in lines_to_add:
                        if line.strip():
                            doc.add_paragraph(line.strip())

                print(f"  Injected Custom Summary ({len(lines_to_add)} lines)")
                changes = True
            else:
                print("  WARNING: Custom Summary heading not found, creating one")
                doc.add_heading("Teams Custom Summary (Loss-Less Extraction)", level=1)
                for line in cs_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                changes = True

    if changes:
        doc.save(str(doc_path))
        print(f"SAVED: {doc_path.name}")
    else:
        print("No changes made.")


if __name__ == "__main__":
    main()
