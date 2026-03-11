import docx
doc_path = '/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/2026-03-10 1300 - SC Internal Deal Desk Project Meeting - Danielle-Bill-Matthew - eid9d3d92a8.docx'
doc = docx.Document(doc_path)
with open('runs/2026-03-02__onenote-org/tmp/clipboard_custom.txt', 'r') as f:
    text = f.read()

doc.add_heading('Teams Recap Extraction (Append-Only)', level=2)
doc.add_paragraph(text)
doc.save(doc_path)
print("Saved custom summary to docx")
