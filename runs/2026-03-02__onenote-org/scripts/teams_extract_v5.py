import sys
import os
import time
from playwright.sync_api import sync_playwright
import docx

def main():
    profile_dir = os.path.abspath('runs/2026-03-02__onenote-org/playwright-profile')
    
    summary_text = ""
    transcript_text = ""

    print("Starting Playwright...")
    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=profile_dir,
            headless=False,
            viewport={"width": 1440, "height": 900}
        )
        page = browser.pages[0] if browser.pages else browser.new_page()
        
        print("Navigating to Teams...")
        page.goto("https://teams.microsoft.com/v2/", wait_until="domcontentloaded")
        
        print("\n" + "="*80)
        print("ACTION REQUIRED:")
        print("1. Find and open the 'SC Internal Deal Desk Project Meeting'.")
        print("2. Stay on the tab for 'AI summary', the bot will extract the data.")
        print("3. Press Enter here when the meeting video page is fully visible.")
        print("="*80 + "\n")
        
        input("Press Enter to continue...")
        
        print("Looking for the 'AI summary' tab...")
        try:
            ai_tab = page.get_by_role("tab", name="AI summary").first
            if not ai_tab.is_visible():
                ai_tab = page.locator('text="AI summary"').first
                
            if ai_tab.is_visible():
                ai_tab.click()
                time.sleep(3)
                
                expand_btn = page.get_by_role("button", name="Expand all").first
                if not expand_btn.is_visible():
                    expand_btn = page.locator('text="Expand all"').first
                
                if expand_btn.is_visible():
                    expand_btn.click()
                    time.sleep(2)
                
                print("Extracting AI summary text...")
                summary_text = page.locator('div[role="tabpanel"]').first.inner_text()
                if not summary_text:
                    summary_text = page.locator("body").inner_text()
            else:
                summary_text = page.locator("body").inner_text() # fallback
        except Exception as e:
            print("Error during AI summary extraction:", e)

        print("Looking for the 'Transcript' tab...")
        try:
            transcript_tab = page.get_by_role("tab", name="Transcript").first
            if not transcript_tab.is_visible():
                transcript_tab = page.locator('text="Transcript"').first
                
            if transcript_tab.is_visible():
                transcript_tab.click()
                time.sleep(5) 
                
                print("Extracting Transcript text...")
                transcript_text = page.locator('div[role="tabpanel"]').first.inner_text()
                if not transcript_text:
                    transcript_text = page.locator("body").inner_text()
            else:
                transcript_text = page.locator("body").inner_text()
        except Exception as e:
             print("Error during Transcript extraction:", e)

        browser.close()

    print("Cleaning up AI Summary text length:", len(summary_text))
    print("Cleaning up Transcript text length:", len(transcript_text))
    
    os.makedirs("runs/2026-03-02__onenote-org/tmp", exist_ok=True)
    with open("runs/2026-03-02__onenote-org/tmp/teams_ai_backup.txt", "w") as f:
        f.write(summary_text)
    with open("runs/2026-03-02__onenote-org/tmp/teams_transcript_backup.txt", "w") as f:
        f.write(transcript_text)

    if not summary_text and not transcript_text:
        print("Failed to extract any text. Exiting before Word document update.")
        return

    print("Locating matching Word document...")
    docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
    target_doc_path = None
    if os.path.exists(docs_dir):
        for file in os.listdir(docs_dir):
            if "SC Internal Deal Desk Project" in file and file.endswith(".docx") and not file.startswith("~$"):
                target_doc_path = os.path.join(docs_dir, file)
                break
    
    if not target_doc_path:
        print(f"Could not find Word document in {docs_dir}")
        return
        
    print(f"Updating doc: {os.path.basename(target_doc_path)}")
    try:
        document = docx.Document(target_doc_path)
        
        already_added = any("Teams AI Summary & Transcript" in p.text for p in document.paragraphs)
        if already_added:
            print("Content already present in document.")
            return
            
        document.add_heading("Teams AI Summary & Transcript", level=1)
        
        document.add_heading("Teams AI Summary", level=2)
        if summary_text:
            for line in summary_text.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Failed to extract AI summary]")
            
        document.add_heading("Full Teams Transcript", level=2)
        if transcript_text:
            document.add_paragraph(transcript_text)
        else:
            document.add_paragraph("[Failed to extract transcript]")
            
        document.save(target_doc_path)
        print("Document successfully updated and saved to OneDrive!")
        
    except Exception as e:
        print(f"Error handling docx: {e}")

if __name__ == "__main__":
    main()
