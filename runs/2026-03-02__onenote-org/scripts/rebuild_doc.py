"""
Rebuild a meeting notes doc: extract metadata from old template, create fresh doc
with new heading structure, re-inject AI summary + custom summary.

Usage:
    python3 rebuild_doc.py <old_doc_path> \
        [--ai-html <html_file>] \
        [--custom-summary <txt_file>] \
        [--output <new_doc_path>]

If --output is not specified, the old doc is backed up with .old suffix and
the new doc replaces it.
"""
import argparse
import sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import injection helpers from the sibling scripts
sys.path.insert(0, str(Path(__file__).parent))
from inject_ai_summary_rich import extract_content_blocks, merge_text_blocks, inject_blocks


def extract_metadata(doc):
    """Extract metadata paragraphs from the old doc (everything before Manual Notes H1)."""
    meta_paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and 'Manual Notes' in p.text:
            break
        meta_paragraphs.append(p)
    return meta_paragraphs


def copy_paragraph(src_para, dst_doc):
    """Copy a paragraph from source doc to destination doc, preserving style and formatting."""
    new_para = dst_doc.add_paragraph()
    new_para.style = dst_doc.styles[src_para.style.name] if src_para.style.name in [s.name for s in dst_doc.styles] else src_para.style
    new_para.alignment = src_para.alignment
    new_para.paragraph_format.space_before = src_para.paragraph_format.space_before
    new_para.paragraph_format.space_after = src_para.paragraph_format.space_after

    # Clear the default empty run
    new_para.clear()

    for run in src_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name

    return new_para


def build_new_doc(old_doc):
    """Create a new doc with metadata from old doc + new heading structure."""
    new_doc = Document()

    # Copy metadata paragraphs (title, time, attendees, description, event ID)
    meta = extract_metadata(old_doc)
    for p in meta:
        copy_paragraph(p, new_doc)

    # New template headings
    new_doc.add_heading("Manual Notes", level=1)

    new_doc.add_heading("Teams AI Summary", level=1)

    new_doc.add_heading("Teams Custom Summary (Loss-Less Extraction)", level=1)

    new_doc.add_heading("Otter Transcript Imports (Append-Only)", level=1)
    new_doc.add_paragraph("Otter transcripts will be appended below.\n")

    return new_doc


def main():
    parser = argparse.ArgumentParser(description="Rebuild meeting notes doc with new template")
    parser.add_argument("old_doc", help="Path to the existing .docx file")
    parser.add_argument("--ai-html", help="Path to saved AI summary HTML file")
    parser.add_argument("--custom-summary", help="Path to saved custom summary text file")
    parser.add_argument("--output", help="Output path (default: overwrite original, backup to .old)")
    args = parser.parse_args()

    old_path = Path(args.old_doc)
    if not old_path.exists():
        print(f"ERROR: {old_path} not found")
        sys.exit(1)

    # Read old doc
    old_doc = Document(str(old_path))

    # Build new doc
    new_doc = build_new_doc(old_doc)

    # Determine output path
    if args.output:
        out_path = Path(args.output)
    else:
        # Backup old doc
        backup_path = old_path.with_suffix('.docx.old')
        if not backup_path.exists():
            import shutil
            shutil.copy2(str(old_path), str(backup_path))
            print(f"Backed up original to {backup_path.name}")
        out_path = old_path

    # Save the new template doc
    new_doc.save(str(out_path))
    print(f"Created new template doc: {out_path.name}")

    # Print structure
    print("\nNew doc structure:")
    for i, p in enumerate(new_doc.paragraphs):
        if p.style.name.startswith('Heading') or p.text.strip():
            print(f"  {i:3d} [{p.style.name}] {p.text[:80]}")

    # Re-inject AI summary if provided
    if args.ai_html:
        html_path = Path(args.ai_html)
        if html_path.exists():
            html = html_path.read_text(encoding='utf-8')
            blocks = extract_content_blocks(html)
            blocks = merge_text_blocks(blocks)
            text_count = sum(1 for b in blocks if b["type"] == "text")
            img_count = sum(1 for b in blocks if b["type"] == "image")
            print(f"\nAI Summary: parsed {text_count} text + {img_count} images")

            # Re-open the saved doc for injection
            doc = Document(str(out_path))
            inject_blocks(doc, str(out_path), blocks, replace=True)
        else:
            print(f"WARNING: AI HTML file not found: {html_path}")

    # Re-inject custom summary if provided
    if args.custom_summary:
        cs_path = Path(args.custom_summary)
        if cs_path.exists():
            cs_text = cs_path.read_text(encoding='utf-8')
            lines = [l for l in cs_text.split('\n') if l.strip()]

            # Re-open doc for injection
            doc = Document(str(out_path))

            # Find "Teams Custom Summary" heading
            idx = None
            for i, p in enumerate(doc.paragraphs):
                if p.style.name.startswith('Heading') and 'Teams Custom Summary' in p.text:
                    idx = i
                    break

            if idx is not None:
                # Find next H1
                insert_before = None
                for i in range(idx + 1, len(doc.paragraphs)):
                    p = doc.paragraphs[i]
                    if p.style.name.startswith('Heading'):
                        lvl = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 99
                        if lvl <= 1:
                            insert_before = i
                            break

                ref_element = doc.paragraphs[insert_before]._element if insert_before else None
                if ref_element is not None:
                    for line in lines:
                        new_p = doc.add_paragraph(line.strip())
                        ref_element.addprevious(new_p._element)
                else:
                    for line in lines:
                        doc.add_paragraph(line.strip())

                doc.save(str(out_path))
                print(f"\nCustom Summary: injected {len(lines)} lines")
            else:
                print("WARNING: 'Teams Custom Summary' heading not found")
        else:
            print(f"WARNING: Custom summary file not found: {cs_path}")

    # Final verification
    final_doc = Document(str(out_path))
    print(f"\nFinal doc structure ({len(final_doc.paragraphs)} paragraphs):")
    for i, p in enumerate(final_doc.paragraphs):
        style = p.style.name
        text = p.text[:80] if p.text else ''
        if style.startswith('Heading'):
            print(f"  {i:3d} [{style}] {text}")
        elif text:
            print(f"  {i:3d} [{style}] {text[:60]}...")


if __name__ == "__main__":
    main()
