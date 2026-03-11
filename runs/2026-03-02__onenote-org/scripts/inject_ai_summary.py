"""
Directly inject AI summary text under "Screenshots / Images" or "AI Summary & Screenshots" 
heading in a meeting notes docx.
"""
import sys
from pathlib import Path
from docx import Document


def main():
    doc_path = sys.argv[1]
    ai_file = sys.argv[2]

    doc = Document(doc_path)
    ai_text = Path(ai_file).read_text(encoding="utf-8")
    lines = [l for l in ai_text.split('\n') if l.strip()]

    # Find the target heading
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            if any(term in p.text for term in ["AI Summary & Screenshots", "Screenshots / Images"]):
                target_idx = i
                break

    if target_idx is None:
        print("WARNING: No AI Summary heading found, appending at end under new heading")
        doc.add_heading("AI Summary & Screenshots", level=2)
        for line in lines:
            doc.add_paragraph(line)
        doc.save(doc_path)
        print(f"SAVED (appended): {Path(doc_path).name}")
        return

    # Find the next heading at same or higher level
    target_level = int(doc.paragraphs[target_idx].style.name.split()[-1])
    insert_before_idx = None
    for i in range(target_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading'):
            lvl = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 99
            if lvl <= target_level:
                insert_before_idx = i
                break

    if insert_before_idx is not None:
        ref_element = doc.paragraphs[insert_before_idx]._element
        for line in reversed(lines):
            new_p = doc.add_paragraph(line)
            ref_element.addprevious(new_p._element)
    else:
        for line in lines:
            doc.add_paragraph(line)

    doc.save(doc_path)
    print(f"SAVED: Injected {len(lines)} lines of AI summary into {Path(doc_path).name}")


if __name__ == "__main__":
    main()
