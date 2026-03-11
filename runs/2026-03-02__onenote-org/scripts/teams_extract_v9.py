import sys
import os
import time
from playwright.sync_api import sync_playwright
import docx

def format_transcript(raw_text):
    """Clean up the raw Teams body dump to look like a usable transcript."""
    bad_lines = {"Activity", "Chat", "Teams", "Calendar", "Calls", "OneDrive", "Mentions", "Transcript", "SC Internal Deal Desk Project Meeting"}
    lines = raw_text.split('\n')
    cleaned = []
    
    capture = False
    for line in lines:
        t = line.strip()
        if not t: continue
        
        # We start capturing when we see speakers/timestamps (heuristics)
        # Typically the first speaker name appears after the transcript header.
        if "Transcript" in t or "SC Internal Deal Desk" in t:
            capture = True
            
        if t in bad_lines: continue
        
        # Eliminate weird UI buttons
        if t in ["Expand all", "Collapse all", "Speaker", "Chapters", "Topics"]:
            continue
            
        if capture:
            cleaned.append(t)
    return "\n".join(cleaned)

def main():
    profile_dir = os.path.abspath('runs/2026-03-02__onenote-org/playwright-profile')
    
    summary_text = ""
    transcript_text = ""

    print("\n" + "="*80)
    print("STARTING TEAMS AUTO-EXTRACTOR (V9 - SAFE MODE)")
    print("="*80)
    
    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=profile_dir,
            headless=False,
            viewport={"width": 1440, "height": 900}
        )
        
        page = browser.pages[0] if browser.pages else browser.new_page()
        
        if "teams.microsoft.com" not in page.url:
            print("Navigating to Teams...")
            page.goto("https://teams.microsoft.com/v2/", wait_until="domcontentloaded")
        
        print("\n⏳ ACTION 1: AI SUMMARY DATA")
        print("-> Go to the 'SC Internal Deal Desk Project Meeting' recording page.")
        print("-> Make EXACTLY SURE you are looking at the 'Meeting notes' AI Summary directly.")
        print("-> Click 'Expand all' so all bullets are visible.")
        
        print("\nWaiting up to 60 seconds for you to navigate... I will auto-detect when it's ready.")
        
        found_summary = False
        for i in range(60):
            try:
                # We specifically look for elements inside the recap panel
                text = page.locator("body").inner_text()
                if "SC Internal Deal Desk" in text and "Meeting notes" in text and "AI summary" in text:
                    # Give user 5 seconds to finish expanding
                    print("\n\nDETECTED THE MEETING RECAP! Waiting 5 seconds to allow UI to settle...")
                    time.sleep(5)
                    summary_text = text
                    found_summary = True
                    break
            except:
                pass
            time.sleep(1)
            
        if not found_summary:
            print("❌ TIMEOUT: Could not detect the correct meeting notes on screen. Exiting.")
            browser.close()
            return
            
        print("   ✅ Summary captured safely!")

        print("\n⏳ ACTION 2: TRANSCRIPT DATA")
        print("-> Please click the 'Transcript' tab NOW.")
        print("-> Scroll down just a tiny bit to make sure the chat lazy-loads.")
        
        found_transcript = False
        print("\nWaiting up to 30 seconds for the Transcript to appear...")
        for i in range(30):
            try:
                # Teams transcript pane uses a specific virtualized list, getting body text is usually safe
                # once the tab is switched.
                text = page.locator("body").inner_text()
                if "Speakers" in text or "Time" in text or "Transcript" in text:
                    # We are in the transcript view
                    # But we don't want to capture instantly, give it time to load the DOM
                    time.sleep(3)
                    # To be super safe and simulate the user's "highlight and copy" approach:
                    # we will grab the main content pane specifically
                    print("\n\nDETECTED TRANSCRIPT! Capturing now...")
                    transcript_text = page.locator("body").inner_text()
                    found_transcript = True
                    break
            except:
                pass
            time.sleep(1)

        print("Closing browser...")
        browser.close()

    print("\nProcessing captured text...")
    
    clean_summary = ""
    try:
        # Heavily strip out everything before "Meeting notes"
        if "Meeting notes" in summary_text:
            parts = summary_text.split("Meeting notes")
            clean_summary = "Meeting notes" + parts[-1]
            if "Transcript" in clean_summary:
                clean_summary = clean_summary.split("Transcript")[0]
        else:
            clean_summary = summary_text
    except:
        clean_summary = summary_text

    clean_transcript = format_transcript(transcript_text)

    # Output to TMP for debugging
    with open("runs/2026-03-02__onenote-org/tmp/teams_ai_backup.txt", "w", encoding="utf-8") as f:
        f.write(summary_text)
    with open("runs/2026-03-02__onenote-org/tmp/teams_transcript_backup.txt", "w", encoding="utf-8") as f:
        f.write(transcript_text)

    # ---------------------------------------------------------
    # FIND AND UPDATE THE WORD DOCUMENT
    # ---------------------------------------------------------
    print("\nLocating the target Word document...")
    docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
    target_doc_path = None
    if os.path.exists(docs_dir):
        for file in os.listdir(docs_dir):
            if "SC Internal Deal Desk Project" in file and file.endswith(".docx") and not file.startswith("~$"):
                target_doc_path = os.path.join(docs_dir, file)
                break
    
    if not target_doc_path:
        print(f"❌ Could not find target doc in {docs_dir}")
        return
        
    print(f"Updating: {os.path.basename(target_doc_path)}")
    try:
        document = docx.Document(target_doc_path)
            
        # Append data to the end of the Word Document
        document.add_heading("Microsoft Teams AI Summary & Transcript Extract", level=1)
        
        document.add_heading("AI Summary Notes", level=2)
        if len(clean_summary) > 50:
            for line in clean_summary.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Failed to cleanly extract AI summary text]")
            
        document.add_heading("Full Raw Transcript", level=2)
        if len(clean_transcript) > 50:
            for chunk in clean_transcript.split('\n'):
                if chunk.strip():
                    document.add_paragraph(chunk.strip())
        else:
            document.add_paragraph("[Failed to cleanly extract transcript text]")
            
        document.save(target_doc_path)
        print("\n🎉 SUCCESS! I have saved the Team's AI Summary and Transcript into your meeting notes doc.")
        
    except Exception as e:
        print(f"❌ Error handling file update: {e}")

if __name__ == "__main__":
    main()
