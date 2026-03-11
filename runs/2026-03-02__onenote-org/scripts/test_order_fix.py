"""
Test: verify that inject_ai_summary_rich.py preserves content order
when injecting into a doc that has headings after the target section.
"""
from pathlib import Path
from docx import Document

# Create a minimal doc with the same heading structure as meeting notes
test_doc_path = "runs/2026-03-02__onenote-org/tmp/test_order_with_headings.docx"
doc = Document()
doc.add_heading("Meeting Title", level=1)
doc.add_heading("Manual Notes", level=1)
doc.add_heading("Screenshots / Images", level=2)
doc.add_paragraph("[placeholder]")
doc.add_heading("Teams Recap Extraction (Append-Only)", level=1)
doc.add_paragraph("[recap placeholder]")
doc.add_heading("Otter Transcript Imports (Append-Only)", level=1)
doc.add_paragraph("[otter placeholder]")
doc.save(test_doc_path)
print(f"Created test doc: {test_doc_path}")

# Now inject using the fixed script
import subprocess, sys
result = subprocess.run(
    [sys.executable, "runs/2026-03-02__onenote-org/scripts/inject_ai_summary_rich.py",
     test_doc_path,
     "--from-file", "runs/2026-03-02__onenote-org/tmp/deal_desk_ai_summary.html"],
    capture_output=True, text=True
)
print(result.stdout)
if result.stderr:
    print("STDERR:", result.stderr)

# Read back and verify order
doc = Document(test_doc_path)
print("\n" + "=" * 80)
print("DOCUMENT CONTENT ORDER (after injection into doc with headings)")
print("=" * 80)
for i, p in enumerate(doc.paragraphs):
    style = p.style.name
    text = p.text[:100] if p.text else ""
    drawings = p._element.findall(
        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
    )
    if drawings:
        print(f"  {i:3d}  [{style:15s}]  [IMAGE x{len(drawings)}]")
    elif text:
        print(f"  {i:3d}  [{style:15s}]  {text}")
    elif style.startswith('Heading'):
        print(f"  {i:3d}  [{style:15s}]  (empty)")
print("=" * 80)
