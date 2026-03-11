import os
import docx

docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
try:
    target_doc_path = [os.path.join(docs_dir, f) for f in os.listdir(docs_dir) if "SC Internal Deal Desk Project" in f and not f.startswith("~$")][0]
    print(f"Cleaning: {target_doc_path}")
    doc = docx.Document(target_doc_path)

    delete_from = -1
    for i, p in enumerate(doc.paragraphs):
        # We look for the exact string we added earlier
        if "Teams AI Summary & Transcript" in p.text:
            delete_from = i
            break

    if delete_from != -1:
        for p in list(doc.paragraphs)[delete_from:]:
            p._element.getparent().remove(p._element)
        doc.save(target_doc_path)
        print("Successfully cleaned old transcript block from doc.")
    else:
        print("Nothing to clean.")
except Exception as e:
    print(f"Error: {e}")
