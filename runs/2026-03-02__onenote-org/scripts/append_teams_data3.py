import json
import base64
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches

target_doc = '/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/2026-03-10 1300 - SC Internal Deal Desk Project Meeting - Danielle-Bill-Matthew - eid9d3d92a8.docx'
doc = Document(target_doc)

img_file = '/Users/matt_bossemeyer/Library/Application Support/Code/User/workspaceStorage/b2c8a6b433c5f01236f3b35f6576baa6/GitHub.copilot-chat/chat-session-resources/21c274bd-9259-4404-9711-93e99e57b32d/call_MHxKQkpEZUxERXYxVXd5c1lPR2g__vscode-1773158419224/content.txt'
txt_file = '/Users/matt_bossemeyer/Library/Application Support/Code/User/workspaceStorage/b2c8a6b433c5f01236f3b35f6576baa6/GitHub.copilot-chat/chat-session-resources/21c274bd-9259-4404-9711-93e99e57b32d/call_MHx3bXlrbnRaN0xOdGxrTG1nMjc__vscode-1773158419222/content.txt'

# Load text
with open(txt_file, 'r') as f:
    text_data_str = f.read()

if text_data_str.startswith('### Result\n'):
    text_data_str = text_data_str[11:]
if '### Ran Playwright code' in text_data_str:
    text_data_str = text_data_str.split('### Ran Playwright code')[0].strip()

text_json = json.loads(text_data_str)
transcript_text = text_json.get('transcript', '')

# Append full transcript
doc.add_heading('Teams Transcript (Full Extracted - Run 2)', level=2)
for paragraph in transcript_text.split('\n'):
    if paragraph.strip():
        doc.add_paragraph(paragraph.strip())

# Load images
with open(img_file, 'r') as f:
    img_data_str = f.read()

if img_data_str.startswith('### Result\n'):
    img_data_str = img_data_str[11:]
if '### Ran Playwright code' in img_data_str:
    img_data_str = img_data_str.split('### Ran Playwright code')[0].strip()

img_json = json.loads(img_data_str)
images = img_json.get('images', [])
ai_notes = img_json.get('aiNotes')

doc.add_heading('Teams AI Summary and Images', level=2)
if ai_notes:
    doc.add_paragraph(ai_notes)

img_count = 0
for img in images:
    if img.get('full_b64'):
        b64str = img['full_b64'].split(',')[1] if ',' in img['full_b64'] else img['full_b64']
        try:
            img_bytes = base64.b64decode(b64str)
            image_stream = BytesIO(img_bytes)
            # Add a bit of space or text
            if img.get('alt'):
                doc.add_paragraph(f"Screenshot / Note: {img['alt']}")
            doc.add_picture(image_stream, width=Inches(6.0))
            img_count += 1
        except Exception as e:
            print(f'Error adding image: {e}')

doc.save(target_doc)
print(f'Successfully added {len(transcript_text)} transcript chars and {img_count} images to doc.')
