import sys
import os
import time
from playwright.sync_api import sync_playwright
import docx

def main():
    profile_dir = os.path.abspath('runs/2026-03-02__onenote-org/playwright-profile')
    
    summary_text = ""
    transcript_text = ""

    print("Starting Playwright Automator...")
    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=profile_dir,
            headless=False,
            viewport={"width": 1440, "height": 900}
        )
        page = browser.pages[0] if browser.pages else browser.new_page()
        
        print("Navigating to Teams...")
        page.goto("https://teams.microsoft.com/v2/", wait_until="load")
        
        print("\n" + "*" * 80)
        print("I have opened the browser!")
        print("1. Find the 'SC Internal Deal Desk Project Meeting' recording.")
        print("2. Open it so that the video and 'AI summary' tab are visible.")
        print("3. CLICK 'Expand all' under the Meeting notes if it exists.")
        print("4. When you perfectly see the notes, come back here and PRESS ENTER.")
        print("*" * 80 + "\n")
        
        # Wait for the user to confirm they have the right screen open
        input("Press ENTER when you are looking at the 'AI summary' tab...")
        
        try:
            print("\nCapturing everything visible on the screen...")
            # Rather than messing with complex selectors, try to grab the literal visible text of the app
            summary_text = page.locator("body").inner_text()
        except Exception as e:
            print("Error grabbing summary:", e)

        print("\nNow, manually click the 'Transcript' tab in the browser.")
        print("Once the Transcript loads and you can see the scrolling text of the speakers:")
        input("Press ENTER here to capture it...")

        try:
            print("Capturing Transcript text...")
            transcript_text = page.locator("body").inner_text()
        except Exception as e:
            print("Error grabbing transcript:", e)

        # Let's take a screenshot of what we saw just in case
        os.makedirs("runs/2026-03-02__onenote-org/tmp", exist_ok=True)
        try:
            page.screenshot(path="runs/2026-03-02__onenote-org/tmp/teams_last_state.png")
        except:
            pass

        browser.close()

    print(f"\nExtraction complete. Summary raw size: {len(summary_text)} | Transcript raw size: {len(transcript_text)}")
    
    # Save the raw dumps so we have them
    with open("runs/2026-03-02__onenote-org/tmp/teams_ai_backup.txt", "w", encoding="utf-8") as f:
        f.write(summary_text)
    with open("runs/2026-03-02__onenote-org/tmp/teams_transcript_backup.txt", "w", encoding="utf-8") as f:
        f.write(transcript_text)

    # Simple cleanup heuristics to remove surrounding UI junk
    # Look for the marker words in the AI summary dump
    clean_summary = ""
    try:
        if "Meeting notes" in summary_text:
            parts = summary_text.split("Meeting notes")
            # Usually the notes are after the header
            clean_summary = parts[-1].strip()
            # Stop if we hit typical UI elements below it
            if "Transcript" in clean_summary:
                clean_summary = clean_summary.split("Transcript")[0]
        else:
            clean_summary = summary_text
    except Exception:
        clean_summary = summary_text

    clean_transcript = ""
    try:
        # Transcript list usually consists of Speaker names + Time + Text
        # The main body dump should have it all. Let's just use the raw dump for now.
        clean_transcript = transcript_text
    except:
        clean_transcript = transcript_text

    print("Locating matching Word document...")
    docs_dir = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/"
    target_doc_path = None
    if os.path.exists(docs_dir):
        for file in os.listdir(docs_dir):
            if "SC Internal Deal Desk Project" in file and file.endswith(".docx") and not file.startswith("~$"):
                target_doc_path = os.path.join(docs_dir, file)
                break
    
    if not target_doc_path:
        print(f"Could not find target doc in {docs_dir}")
        return
        
    print(f"Updating doc: {os.path.basename(target_doc_path)}")
    try:
        document = docx.Document(target_doc_path)
        
        document.add_heading("Microsoft Teams AI Summary & Transcript", level=1)
        
        document.add_heading("Teams AI Summary", level=2)
        if clean_summary and len(clean_summary) > 20:
            for line in clean_summary.split('\n'):
                if line.strip():
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Extracted data too short or missing]")
            
        document.add_heading("Teams Raw Transcript", level=2)
        if clean_transcript and len(clean_transcript) > 20:
            for line in clean_transcript.split('\n'):
                # filter out some obvious top-level teams UI boilerplate if we can
                if line.strip() and not line.strip() in ['Activity', 'Chat', 'Teams', 'Calendar', 'Calls', 'OneDrive']:
                    document.add_paragraph(line.strip())
        else:
            document.add_paragraph("[Extracted data too short or missing]")
            
        document.save(target_doc_path)
        print("\nSUCCESS! I have injected the Team's AI Summary and Transcript into your meeting notes doc.")
        
    except Exception as e:
        print(f"Error handling docx: {e}")

if __name__ == "__main__":
    main()
