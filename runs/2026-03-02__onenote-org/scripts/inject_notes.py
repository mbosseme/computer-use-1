import os
import docx

doc_path = "/Users/matt_bossemeyer/Library/CloudStorage/OneDrive-Premier,Inc/Notetaking/Meeting Notes/2026/03/10/2026-03-10 1300 - SC Internal Deal Desk Project Meeting - Danielle-Bill-Matthew - eid9d3d92a8.docx"
doc = docx.Document(doc_path)

recap_text = """### EXHAUSTIVE TOPIC EXTRACTION
- **Projects & Initiatives**
- At the start of the year, Premier identified 10 focus projects; 3 are supply chain-related, with Bruce leading or co-leading two. These two projects are expected to directly impact the future operation of the deal desk.
- The deal desk build-out is envisioned in three phases: immediate setup for impact, iterative improvement as projects deliver, and future integration with AI/technology tools.
- Growth project includes an "opportunity heat map" analyzing contract uptake and off-contract spend. Estimated coverage of Premier contracts is about 40% of members' non-labor expense.
- The heat map identifies categories with low uptake and explores reasons (delivery, packaging, contract gaps). Nearly equal spend is on and off contract.
- Second part of growth project: identifies member spend in categories not covered by Premier contracts (e.g., Epic/EMRs unlikely to be GPO-contracted).
- Admin fee project analyzes fee rates across four business units, scenario modeling increases, and aims to deliver target ranges by category and a heat map of gaps by end of summer.
- Palantir and other tools are expected to go live in the fall, creating a "scoreboard" for contract value.
- Value Assurance team uses Healthcare IQ benchmarks, with percentiles (closer to 100 is better per HCIQ, but Premier historically used lower-is-better).
- Typical contract execution timeline is 17 months.
- National portfolio pricing is currently at the 46th percentile; targeting 75th percentile would yield ~11% savings.
- Value Assurance team is six people; HCIQ benchmarks are new to sourcing groups.
- Categories discussed for analysis: exam gloves, gastro/endo, CRM, pick and midline access products, contrast media ultrasound, specialty women's health, incontinence.
- Incontinence category: five suppliers on national, but Medline is representative; 17% gap on Ascend, 32% gap on national.
- Data resides in ABI platform, Passport, shared Excel lists (clinical), and possibly Monday.com (non-clinical/services).
- Key contacts for data visibility: Kyle MacKinnon (clinical), Rachel Rollins (non-clinical/services).

- **Metrics & KPIs**
- Contract attributes for deal desk review: price change, admin fee rate change, supplier count change, coverage percentage.
- Admin fee scenario modeling: increases by 0.1% or 0.5% considered.
- National contract pricing: current average at 46th percentile, target is 75th percentile for 11% improvement.
- Surpass program: lowest in market; Ascend Drive: 10th percentile low; National top tier: 25th percentile low (Premier convention).
- Incontinence: 17% gap on Ascend, 32% gap on national.

- **System Names & Tools**
- ABI platform, Passport, Monday.com, Palantir, Healthcare IQ, Value Assurance scorecard.

### STRATEGIC INSIGHTS & CONSTRAINTS
- **Core Problems**
- National portfolio contracts lack competitive value; need to "stop the bleeding" and ensure all contracts post-review are valuable.
- Aggregation groups leverage national portfolio, sometimes encroaching on committed programs.
- Sourcing teams need actionable standards and benchmarks to avoid manual review of all contracts; focus should be on exceptions.
- Rebates complicate benchmarking; CRM is rebate-heavy, and rebate tracking will be integrated with Palantir.
- Multiple goals can be counterintuitive: Premier values admin fees, members value price and coverage.

### SPEAKER ATTRIBUTION & FRICTION POINTS
- **Matt Shimshock**: Pushed for clarity on deal desk's purpose: value vs. growth. Suggested scenario modeling and optimization for contract categories.
- **Ken**: Emphasized fewer suppliers yield better prices and market share shifts.
- **Friction Points & Debate**: Debate on whether deal desk should focus on growth/admin fees or contract value; consensus formed around prioritizing value.
"""

otter_index = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Teams Recap Extraction (Append-Only)"):
        otter_index = i
        break

if otter_index != -1:
    p = doc.paragraphs[otter_index]
    doc.paragraphs[otter_index+1].insert_paragraph_before("--- Playwright MCP Extracted Summary ---")
    doc.paragraphs[otter_index+1].insert_paragraph_before(recap_text.strip())
    inserted = True
else:
    doc.add_paragraph("--- Playwright MCP Extracted Summary ---")
    doc.add_paragraph(recap_text.strip())
    inserted = True

doc.save(doc_path)
print("Insertion completed successfully:", inserted)